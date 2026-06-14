import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

paras = json.load(open('/tmp/paras.json'))

def load_chapter(path):
    raw = open(path).read().split('\n')
    lines = [l.strip() for l in raw if l.strip()!='']
    out = []
    out.append(('h1', lines[0]))           # CHAPTER X: Title
    for l in lines[1:]:
        out.append(('body', l))
    return out

ch4  = load_chapter('/home/user/test/chapter-04-the-wrestler-highlights.txt')
ch10 = load_chapter('/home/user/test/chapter-10-west-palm-beach.txt')
ch11 = load_chapter('/home/user/test/chapter-11-the-uhaul.txt')
ch12 = load_chapter('/home/user/test/chapter-12-january-ninth.txt')

MAJOR = {
 'PROLOGUE: THE CARDINAL','CHAPTER 1: Mexico City, 1982 — The Escape',
 'CHAPTER 2: New Jersey — Survival',"CHAPTER 3: The Man Who Didn't Have To",
 'CHAPTER 4: The Wrestler',"CHAPTER 5: My Brother's Keeper","CHAPTER 6: Lou's World",
 'CHAPTER 7: Oxford','CHAPTER 8: The Performance Trap','HAPTER 9: Checking Boxes',
 'CHAPTER 13: The Quiet Years','CHAPTER 14: The Line','CHAPTER 15: The Salmon Palace',
 'CHAPTER 16: The Offer','CHAPTER 17: They Gave Me a Chance','CHAPTER 18: Getting Square',
 'CHAPTER 19: Gulf Breeze','CHAPTER 20: The People You Meet When You Stop Drinking',
 'CHAPTER 21: Hope','CHAPTER 23: The Pictures Never Tell the Whole Story',
 'CHAPTER 24: Blessings','CHAPTER 25: The Weight That Stays','CHAPTER 26: The Detour',
 'CHAPTER 27: Five Weeks','THE SUIT','THE TABLE','WHAT GRIEF TAUGHT THE JOB',
 'THE ROUND TABLE',"EPILOGUE: The Cardinal's Promise","WHAT HE DIDN'T OWE",
 'WHO STAYS. WHO GOES.',
}

SKIP = {43,80,283,303,385,139}           # timestamps + duplicate Baha Mar paragraph
SKIP |= set(range(352,373))              # duplicate CHAPTER 21 block (352-372)
CH4_BLOCK = set(range(81,124))           # replaced by condensed version

# Build ordered (kind, text) stream
stream = []
for i, t in enumerate(paras):
    s = t.strip()
    if i == 81:                          # inject condensed Chapter 4 in place of old block
        stream.extend(ch4)
    if i in CH4_BLOCK:
        continue
    if i in SKIP:
        continue
    if i == 237:                         # before CHAPTER 13 -> insert missing 10,11,12
        stream.extend(ch10); stream.extend(ch11); stream.extend(ch12)
    if i == 373:                         # before "Wonderwoman" -> add missing Ch22 heading
        stream.append(('h1','CHAPTER 22: Wonderwoman'))
    if s == '':
        continue
    if i == 0:
        stream.append(('title', s)); continue
    if i in (1,2,3):
        stream.append(('subtitle', s)); continue
    if s in MAJOR:
        stream.append(('h1', 'CHAPTER 9: Checking Boxes' if s=='HAPTER 9: Checking Boxes' else s)); continue
    # subsection heading heuristic
    if len(s) < 58 and not s.endswith('.') and not s.endswith('”') and len(s.split()) <= 9 and s[0].isupper() and s != '— R.B.':
        stream.append(('h2', s)); continue
    stream.append(('body', s))

# Render
doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Garamond'; normal.font.size = Pt(12)
doc.styles['Normal'].paragraph_format.space_after = Pt(8)

for kind, text in stream:
    if kind == 'title':
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = True; r.font.size = Pt(26)
    elif kind == 'subtitle':
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.italic = True; r.font.size = Pt(13)
    elif kind == 'h1':
        doc.add_page_break()
        p = doc.add_paragraph(); p.space_before = Pt(0)
        r = p.add_run(text); r.bold = True; r.font.size = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    elif kind == 'h2':
        p = doc.add_paragraph()
        r = p.add_run(text); r.bold = True; r.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
    else:
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.15

doc.save('/home/user/test/The Cardinals Promise - draft 4.13.2026.docx')

# stats
words = sum(len(t.split()) for k,t in stream if k=='body')
print('paragraphs rendered:', len(stream))
print('approx body word count:', words)
